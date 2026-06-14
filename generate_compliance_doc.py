"""Genera el documento de compliance FSE+ / RGPD en formato .docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Estilos globales ──────────────────────────────────────────────────────────

style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def set_heading(text, level, color=None):
    h = doc.add_heading(text, level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return h

def add_paragraph(text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    return p

def add_code_block(code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Fondo gris claro via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Cabecera
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Filas
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = cell_text
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

# ── Portada ───────────────────────────────────────────────────────────────────

doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('DOCUMENTACIÓN TÉCNICA DE COMPLIANCE')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run('PLATAFORMA SOCIALPAY MVP')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()

meta_lines = [
    ('Clasificación:', 'Uso Oficial — Up Spain / Auditorías FSE+'),
    ('Versión:', '1.0 — Junio 2026'),
    ('Elaborado por:', 'Departamento de Tecnología'),
    ('Marco normativo:', 'Reglamento FSE+ 2021/1057, RGPD (UE) 2016/679'),
]
for label, value in meta_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + ' ')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.size = Pt(11)

doc.add_page_break()

# ── APARTADO 1 ────────────────────────────────────────────────────────────────

set_heading('APARTADO 1: ARQUITECTURA Y AUDITABILIDAD FSE+', level=1, color=(0x1F, 0x38, 0x64))

# 1.1
set_heading('1.1 Entornos Aislados y Trazabilidad del Ciclo de Vida', level=2)

add_paragraph(
    'La plataforma SocialPay MVP opera sobre una arquitectura de tres entornos completamente aislados, '
    'gestionados a través del servicio de despliegue en la nube Railway. Esta separación garantiza que '
    'ningún dato de prueba o desarrollo contamine el registro oficial de transacciones FSE+.'
)

add_table(
    headers=['Entorno', 'Propósito', 'Base de Datos', 'URL de Acceso'],
    rows=[
        ['Local', 'Desarrollo y pruebas de código', 'SQLite (socialpay_dev.db)', 'localhost:8000'],
        ['Staging', 'Validación pre-producción / QA', 'PostgreSQL (Railway, rama staging)', 'socialpay-staging.up.railway.app'],
        ['Producción', 'Operación real FSE+', 'PostgreSQL (Railway, rama main)', 'socialpay-production.up.railway.app'],
    ],
    col_widths=[1.1, 1.8, 2.2, 2.4],
)

add_paragraph(
    'El módulo app/database.py implementa la detección de entorno mediante la variable de sistema DATABASE_URL:'
)
add_code_block(
    'DATABASE_URL = os.environ.get("DATABASE_URL", "")\n'
    'if DATABASE_URL:\n'
    '    if DATABASE_URL.startswith("postgres://"):\n'
    '        DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql://", 1)\n'
    'else:\n'
    '    DATABASE_URL = "sqlite:///./socialpay_dev.db"  # Fallback local'
)
add_paragraph(
    'Este mecanismo garantiza que el motor de base de datos se configure de forma completamente automática '
    'según el entorno, sin posibilidad de error de configuración manual que pudiera mezclar datos entre '
    'entornos. Los entornos de Staging y Producción utilizan bases de datos PostgreSQL alojadas en instancias '
    'Railway independientes, con credenciales separadas y sin acceso cruzado entre ellas.'
)

# 1.2
set_heading('1.2 Registro Inmutable de Transacciones — Modelo AuditoriaTransaccion', level=2)

add_paragraph(
    'La tabla auditoria_transacciones_sqlalchemy constituye el registro Append-Only de todas las transacciones '
    'FSE+ procesadas por la plataforma. Su diseño garantiza la inmutabilidad por construcción: no existe en el '
    'backend ningún endpoint DELETE, PUT ni PATCH que opere sobre esta tabla.'
)
add_paragraph('Definición del modelo (app/models.py):', bold=True)
add_code_block(
    'class AuditoriaTransaccion(Base):\n'
    "    __tablename__ = 'auditoria_transacciones_sqlalchemy'\n\n"
    '    id             = Column(GUID(), primary_key=True, default=uuid.uuid4)\n'
    "    usuario_uuid   = Column(GUID(), ForeignKey('usuarios.id', ondelete='SET NULL'), nullable=True)\n"
    '    supermercado_id = Column(String, nullable=False, index=True)\n'
    '    total          = Column(Numeric(10, 2), nullable=False)\n'
    '    timestamp      = Column(DateTime, default=datetime.utcnow, nullable=False)\n'
    "    estado         = Column(String, nullable=False)  # 'APPROVED', 'DISCREPANCY', 'REJECTED'"
)

add_paragraph('Garantías de inmutabilidad:', bold=True)

bullets = [
    ('Clave primaria UUID: ', 'Cada transacción recibe un identificador único generado por uuid.uuid4() en el momento de su creación. Este identificador no puede reutilizarse ni modificarse.'),
    ('Timestamp de servidor: ', 'El campo timestamp se asigna con datetime.utcnow en el servidor en el momento de la escritura, siendo ajeno a cualquier manipulación por parte del cliente.'),
    ('Ausencia de endpoints de modificación: ', 'La API no expone ninguna ruta que permita actualizar o eliminar registros de auditoría. Los registros se crean exclusivamente en el flujo de validación de tickets (POST /upload-ticket) y pueden consultarse en modo solo lectura por los roles admin y upspain vía GET /api/admin/audit-logs.'),
    ('Estado de transacción declarativo: ', 'El campo estado registra el resultado de la validación en el momento de la transacción (APPROVED, DISCREPANCY, REJECTED) y no se modifica posteriormente, preservando la integridad del historial de auditoría.'),
]
for bold_part, normal_part in bullets:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(bold_part)
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(normal_part)
    r2.font.size = Pt(11)

doc.add_paragraph()
add_paragraph('Flujo de escritura controlado (app/routers/beneficiario.py):', bold=True)
add_code_block(
    'new_audit = ATModel(\n'
    '    usuario_uuid=user.id,\n'
    '    supermercado_id=supermarket,\n'
    '    total=Decimal(str(cart_total)),\n'
    '    estado="APPROVED",\n'
    ')\n'
    'db.add(new_audit)\n'
    'db.commit()'
)
add_paragraph(
    'La única escritura en la tabla de auditoría se produce tras la validación positiva del ticket OCR por el '
    'módulo TicketValidator. No existe ningún código en el backend que ejecute db.delete(), UPDATE ni '
    'db.query(ATModel).delete() sobre esta entidad.'
)

# 1.3
set_heading('1.3 Control de Concurrencia Presupuestaria — with_for_update()', level=2)

add_paragraph(
    'Para garantizar que el presupuesto FSE+ asignado por Up Spain a los gestores nunca sufra sobregasto ni '
    'duplicidad en escenarios de alta concurrencia (múltiples transacciones simultáneas del mismo gestor), '
    'la plataforma implementa bloqueo pesimista a nivel de base de datos mediante la cláusula SQL SELECT FOR UPDATE.'
)
add_paragraph('Modelo de presupuesto (app/models.py):', bold=True)
add_code_block(
    'class AsignacionFondosGestor(Base):\n'
    "    __tablename__ = 'asignacion_fondos_gestor'\n\n"
    '    id                    = Column(GUID(), primary_key=True, default=uuid.uuid4)\n'
    "    gestor_id             = Column(GUID(), ForeignKey('usuarios.id', ondelete='CASCADE'))\n"
    '    codigo_proyecto_fse   = Column(String, nullable=False, index=True)\n'
    '    presupuesto_total     = Column(Numeric(10, 2), nullable=False)\n'
    '    presupuesto_consumido = Column(Numeric(10, 2), default=0.00, nullable=False)\n'
    '    tasa_cofinanciacion   = Column(Numeric(5, 2), nullable=False)\n\n'
    '    __table_args__ = (\n'
    '        CheckConstraint(\n'
    "            'presupuesto_consumido <= presupuesto_total',\n"
    "            name='check_presupuesto_consumido_limit'\n"
    '        ),\n'
    '    )'
)

add_paragraph('Doble capa de protección anti-sobregasto:', bold=True)

bullets2 = [
    ('Capa 1 — Restricción a nivel de base de datos (CheckConstraint): ',
     'La restricción check_presupuesto_consumido_limit es una constraint SQL que la propia base de datos '
     'PostgreSQL aplica en cada operación INSERT o UPDATE. Ninguna transacción puede llevar presupuesto_consumido '
     'por encima de presupuesto_total, incluso si se ejecutan operaciones directas sobre la base de datos sin '
     'pasar por la capa de aplicación. Esta es la última línea de defensa del sistema.'),
    ('Capa 2 — Bloqueo pesimista a nivel de aplicación (with_for_update()): ',
     'Antes de registrar un nuevo beneficiario y asignarle saldo, la aplicación adquiere un bloqueo exclusivo '
     'sobre el registro de asignación del gestor (app/routers/admin.py). Esta instrucción genera el comando SQL '
     'SELECT ... FOR UPDATE, que bloquea la fila en PostgreSQL durante la duración de la transacción. Cualquier '
     'otra solicitud concurrente que intente acceder al mismo registro de asignación quedará bloqueada hasta que '
     'la transacción en curso se complete o revierta, eliminando la condición de carrera (race condition) que '
     'podría producir sobregastos en entornos de carga alta.'),
]
for bold_part, normal_part in bullets2:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(bold_part)
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(normal_part)
    r2.font.size = Pt(11)

doc.add_paragraph()
add_code_block(
    'allocation = tx_session.query(AFGModel).filter(\n'
    '    AFGModel.gestor_id == gestor_id,\n'
    '    AFGModel.codigo_proyecto_fse == codigo_proyecto_fse,\n'
    ').with_for_update().first()'
)

add_paragraph('Proceso atómico de alta de beneficiario:', bold=True)
steps = [
    'Se inicia una transacción de base de datos explícita (with tx_session.begin()).',
    'Se bloquea el registro AsignacionFondosGestor con with_for_update().',
    'Se verifica que el saldo total ya asignado a beneficiarios no supera el presupuesto autorizado.',
    'Si la verificación es correcta, se crea el beneficiario y se actualiza el consumo.',
    'Si la verificación falla, la transacción se revierte sin ninguna escritura parcial.',
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(step).font.size = Pt(11)

doc.add_paragraph()
add_paragraph(
    'Esta arquitectura garantiza que el sistema es ACID-compliant (Atomicidad, Consistencia, Aislamiento, '
    'Durabilidad) para todas las operaciones que afectan al presupuesto FSE+.'
)

doc.add_page_break()

# ── APARTADO 2 ────────────────────────────────────────────────────────────────

set_heading('APARTADO 2: CUMPLIMIENTO ESTRICTO DEL RGPD', level=1, color=(0x1F, 0x38, 0x64))

# 2.1
set_heading('2.1 Principio de Privacidad por Diseño (Privacy by Design)', level=2)

add_paragraph(
    'La plataforma SocialPay MVP ha sido diseñada desde su concepción bajo el principio de privacidad por '
    'diseño y por defecto establecido en el artículo 25 del Reglamento General de Protección de Datos '
    '(RGPD, UE 2016/679). Este principio exige que las medidas técnicas que garantizan la protección de '
    'datos sean parte intrínseca del diseño del sistema, no añadidos posteriores.'
)
add_paragraph(
    'La manifestación técnica fundamental de este principio en SocialPay es la disociación estructural entre '
    'identidad real y operación financiera: el sistema procesa transacciones económicas de beneficiarios '
    'vulnerables sin necesitar conocer, almacenar ni exponer en ningún momento quiénes son esas personas '
    'en el mundo real.'
)

# 2.2
set_heading('2.2 Pseudonimización de Beneficiarios mediante Token Anónimo', level=2)

add_paragraph(
    'El colectivo de usuarios más sensible de la plataforma —los beneficiarios del programa de ayudas FSE+— '
    'no es representado en la base de datos mediante ningún atributo de identificación personal. El modelo '
    'Usuario (app/models.py) para beneficiarios contiene exclusivamente:'
)
add_code_block(
    'class Usuario(Base):\n'
    "    __tablename__ = 'usuarios'\n\n"
    '    id               = Column(GUID(), primary_key=True, default=uuid.uuid4)\n'
    '    token_anonimo    = Column(String, unique=True, index=True, nullable=False)\n'
    '    saldo_disponible = Column(Numeric(10, 2), default=0.00, nullable=False)\n'
    "    rol              = Column(String, nullable=False, default='beneficiario')\n"
    "    gestor_uuid      = Column(GUID(), ForeignKey('usuarios.id', ondelete='SET NULL'), nullable=True)\n"
    '    codigo_proyecto_fse = Column(String, nullable=True)\n\n'
    '    # Solo para personal técnico (roles: admin, upspain, gestor, supermercado)\n'
    '    email            = Column(String, unique=True, index=True, nullable=True)\n'
    '    hashed_password  = Column(String, nullable=True)'
)

add_paragraph(
    'Los campos email y hashed_password son nullable=True y permanecen nulos para todos los beneficiarios. '
    'El sistema no solicita, no almacena ni procesa en ningún momento ninguno de los siguientes atributos '
    'para este colectivo:'
)

add_table(
    headers=['Dato de Carácter Personal (PII)', 'Almacenado en BD', 'Requerido para operar'],
    rows=[
        ['Nombre y apellidos', 'NO', 'NO'],
        ['DNI / NIE / Pasaporte', 'NO', 'NO'],
        ['Dirección postal', 'NO', 'NO'],
        ['Correo electrónico', 'NO', 'NO'],
        ['Número de teléfono', 'NO', 'NO'],
        ['Fecha de nacimiento', 'NO', 'NO'],
        ['Datos bancarios', 'NO', 'NO'],
    ],
    col_widths=[3.2, 1.5, 1.8],
)

add_paragraph(
    'El único identificador del beneficiario en todo el sistema es el token_anonimo: una cadena opaca '
    '(generada por el gestor social responsable) que no contiene ni puede inferir ningún dato personal. '
    'Este token actúa como pseudónimo en el sentido técnico del artículo 4(5) del RGPD: los datos '
    'procesados bajo este identificador no pueden atribuirse a una persona física sin información adicional '
    'que reside exclusivamente fuera del sistema.'
)

# 2.3
set_heading('2.3 Flujo de Autenticación y Operación sin Exposición de PII', level=2)

add_paragraph(
    'El acceso del beneficiario a la plataforma sigue el siguiente flujo, en el que en ningún momento '
    'se solicita ni transmite información personal identificable:'
)
flow_steps = [
    ('Acceso: ', 'El beneficiario introduce su token_anonimo en el formulario de la aplicación móvil.'),
    ('Verificación: ', 'El sistema comprueba si ese token existe en la tabla usuarios (WHERE token_anonimo = ?). No se accede a ningún otro atributo personal en esta verificación.'),
    ('Sesión: ', 'Tras la verificación exitosa, la plataforma establece una cookie de sesión httponly que contiene únicamente el propio token_anonimo. No se almacena ningún JWT con claims personales ni se transmite ningún dato identificable.'),
    ('Operación: ', 'Todas las operaciones posteriores (escaneo de productos, validación de tickets, consulta de saldo) se identifican mediante el token de sesión. El UUID interno se utiliza como referencia de integridad referencial en la base de datos, pero nunca se expone a través de la interfaz de usuario.'),
    ('Registro de auditoría: ', 'Al registrar una transacción en AuditoriaTransaccion, el campo usuario_uuid almacena el UUID interno del beneficiario. Este UUID no es un dato personal por sí mismo: no contiene nombre, DNI ni ningún atributo identificador. La correlación con una persona real solo sería posible con acceso físico a la base de datos del sistema del gestor social externo, que mantiene la tabla de correspondencia token ↔ identidad real, completamente fuera del alcance de SocialPay.'),
]
for i, (bold_part, normal_part) in enumerate(flow_steps, 1):
    p = doc.add_paragraph(style='List Number')
    r1 = p.add_run(bold_part)
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(normal_part)
    r2.font.size = Pt(11)

doc.add_paragraph()

# 2.4
set_heading('2.4 Derecho de Supresión y Preservación del Registro Contable FSE+', level=2)

add_paragraph(
    'El RGPD reconoce en su artículo 17 el derecho al olvido (right to erasure): el interesado puede '
    'solicitar la supresión de sus datos personales. Sin embargo, el artículo 17(3)(b) establece una '
    'excepción cuando el tratamiento es necesario para el cumplimiento de una obligación legal o para la '
    'realización de una misión en interés público, como es el caso de la justificación de fondos europeos FSE+.'
)
add_paragraph(
    'La plataforma SocialPay resuelve esta tensión normativa mediante la directiva ondelete=\'SET NULL\' '
    'en la clave foránea de AuditoriaTransaccion:'
)
add_code_block(
    'usuario_uuid = Column(\n'
    '    GUID(),\n'
    "    ForeignKey('usuarios.id', ondelete='SET NULL'),\n"
    '    nullable=True\n'
    ')'
)
add_paragraph(
    'Cuando se elimina un registro de la tabla usuarios (ejercicio del derecho de supresión), el campo '
    'usuario_uuid en todos los registros de auditoría asociados se establece automáticamente a NULL mediante '
    'una operación a nivel de base de datos, sin intervención de la capa de aplicación.'
)
add_paragraph('Resultado funcional:', bold=True)

results_bullets = [
    'El registro de auditoría FSE+ —imprescindible para las justificaciones financieras ante la Comisión Europea y el Ayuntamiento— se preserva íntegramente con su importe, supermercado, timestamp y estado.',
    'La vinculación con la identidad del beneficiario queda rota de forma permanente e irreversible: la fila de auditoría pasa a tener usuario_uuid = NULL, siendo contablemente válida pero imposible de atribuir a ninguna persona física.',
]
for bullet in results_bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(bullet).font.size = Pt(11)

doc.add_paragraph()
add_paragraph('Este mecanismo cumple simultáneamente con:', bold=True)
compliance_bullets = [
    'Artículo 17 RGPD: los datos personales del beneficiario son eliminados.',
    'Artículo 5(1)(e) RGPD: los datos no se conservan más tiempo del necesario para los fines del tratamiento.',
    'Reglamento FSE+ 2021/1057: el historial de transacciones permanece íntegro para auditoría.',
]
for bullet in compliance_bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(bullet).font.size = Pt(11)

doc.add_paragraph()

# 2.5
set_heading('2.5 Arquitectura de Roles y Principio de Mínimo Privilegio', level=2)

add_paragraph(
    'Solo el personal técnico con responsabilidades operativas en la plataforma (roles admin, upspain, '
    'gestor, supermercado) almacena datos de identificación personal estrictamente necesarios para su '
    'función: email, contraseña cifrada con PBKDF2-SHA256, y datos institucionales como nombre_institucion, '
    'cif, direccion y responsable. Este personal ha firmado los correspondientes contratos de '
    'confidencialidad y tratamiento de datos.'
)
add_paragraph(
    'El acceso a los endpoints de consulta del registro de auditoría está restringido por control de '
    'acceso basado en roles (RBAC) implementado en app/deps.py:'
)

access_bullets = [
    'GET /api/admin/audit-logs: exclusivo para roles admin.',
    'GET /api/upspain/dashboard-data (incluye historial de transacciones): exclusivo para rol upspain.',
]
for bullet in access_bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(bullet).font.size = Pt(11)

doc.add_paragraph()
add_paragraph(
    'Los gestores y supermercados no tienen acceso al historial de transacciones cruzado entre beneficiarios, '
    'únicamente a los datos operativos propios de su ámbito de gestión.',
    italic=True
)

# ── Pie de documento ──────────────────────────────────────────────────────────

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Fin del documento —')
run.italic = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Guardar ───────────────────────────────────────────────────────────────────

output_path = 'SOCIALPAY_Compliance_FSE_RGPD_v1.0.docx'
doc.save(output_path)
print(f"Documento generado: {output_path}")
